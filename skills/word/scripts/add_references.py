"""
add_references.py — Generate formatted reference lists in GB/T 7714 or APA style.
Reads reference data from JSON and appends to .docx.
"""

import argparse
import json
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from format_utils import (
    set_run_font, set_paragraph_format, add_heading_paragraph,
    CN_BODY_FONT, CN_HEADING_FONT, EN_FONT
)


def format_gbt7714(ref, index=1):
    """
    Format a reference in GB/T 7714-2015 style.
    
    Supported types: journal, book, conference, thesis, online, newspaper, report
    """
    ref_type = ref.get('type', 'journal')
    authors = ref.get('authors', '')
    year = ref.get('year', '')
    title = ref.get('title', '')
    
    if ref_type == 'journal':
        journal = ref.get('journal', '')
        volume = ref.get('volume', '')
        issue = ref.get('issue', '')
        pages = ref.get('pages', '')
        vol_str = f",{volume}" if volume else ""
        issue_str = f"({issue})" if issue else ""
        pages_str = f":{pages}" if pages else ""
        return f"[{index}] {authors}. {title}[J]. {journal}{vol_str}{issue_str}{pages_str}."
    
    elif ref_type == 'book':
        publisher = ref.get('publisher', '')
        city = ref.get('city', '')
        pages = ref.get('pages', '')
        loc = f"{city}: " if city else ""
        pages_str = f": {pages}" if pages else ""
        return f"[{index}] {authors}. {title}[M]. {loc}{publisher},{year}{pages_str}."
    
    elif ref_type == 'conference':
        conference = ref.get('conference', '')
        city = ref.get('city', '')
        pages = ref.get('pages', '')
        loc = f". {city}" if city else ""
        pages_str = f": {pages}" if pages else ""
        return f"[{index}] {authors}. {title}[C]//{conference}{loc},{year}{pages_str}."
    
    elif ref_type == 'thesis':
        institution = ref.get('institution', '')
        degree = ref.get('degree', '硕士')
        city = ref.get('city', '')
        loc = f"{city}: " if city else ""
        return f"[{index}] {authors}. {title}[D]. {loc}{institution},{year}."
    
    elif ref_type == 'online':
        url = ref.get('url', '')
        access_date = ref.get('access_date', '')
        date_str = f"[{access_date}]" if access_date else ""
        return f"[{index}] {authors}. {title}[EB/OL]. {url},{year}{date_str}."
    
    elif ref_type == 'newspaper':
        newspaper = ref.get('newspaper', '')
        date = ref.get('date', year)
        page = ref.get('page', '')
        page_str = f"({page})" if page else ""
        return f"[{index}] {authors}. {title}[N]. {newspaper},{date}{page_str}."
    
    elif ref_type == 'report':
        institution = ref.get('institution', '')
        city = ref.get('city', '')
        loc = f"{city}: " if city else ""
        return f"[{index}] {authors}. {title}[R]. {loc}{institution},{year}."
    
    else:
        return f"[{index}] {authors}. {title}. {year}."


def format_apa(ref, index=None):
    """
    Format a reference in APA 7th edition style.
    """
    ref_type = ref.get('type', 'journal')
    authors = ref.get('authors', '')
    year = ref.get('year', '')
    title = ref.get('title', '')
    
    if ref_type == 'journal':
        journal = ref.get('journal', '')
        volume = ref.get('volume', '')
        issue = ref.get('issue', '')
        pages = ref.get('pages', '')
        vol_str = f", {volume}" if volume else ""
        issue_str = f"({issue})" if issue else ""
        pages_str = f", {pages}" if pages else ""
        # APA: Author(s). (Year). Title. Journal, Volume(Issue), Pages.
        return f"{authors} ({year}). {title}. *{journal}*{vol_str}{issue_str}{pages_str}."
    
    elif ref_type == 'book':
        publisher = ref.get('publisher', '')
        city = ref.get('city', '')
        return f"{authors} ({year}). *{title}*. {publisher}."
    
    elif ref_type == 'conference':
        conference = ref.get('conference', '')
        pages = ref.get('pages', '')
        pages_str = f" (pp. {pages})" if pages else ""
        return f"{authors} ({year}). {title}. In *{conference}*{pages_str}."
    
    elif ref_type == 'thesis':
        institution = ref.get('institution', '')
        degree = ref.get('degree', 'Master')
        return f"{authors} ({year}). *{title}* [{degree}'s thesis, {institution}]."
    
    elif ref_type == 'online':
        url = ref.get('url', '')
        return f"{authors} ({year}). {title}. {url}"
    
    else:
        return f"{authors} ({year}). {title}."


def add_reference_list(doc, references, style='gbt7714', heading='参考文献'):
    """
    Add a formatted reference list to the document.
    
    Args:
        doc: Document object
        references: List of reference dicts
        style: 'gbt7714' or 'apa'
        heading: Section heading text
    
    Returns:
        List of formatted reference strings
    """
    # Add heading
    add_heading_paragraph(doc, heading, level=1)
    
    formatted = []
    for i, ref in enumerate(references, 1):
        if style == 'gbt7714':
            text = format_gbt7714(ref, index=i)
        else:
            text = format_apa(ref)
        formatted.append(text)
        
        # Add as paragraph
        p = doc.add_paragraph()
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing=1.5,
            space_before=Pt(0),
            space_after=Pt(0),
            left_indent=Cm(0.74),
            first_line_indent=Cm(-0.74)  # Hanging indent
        )
        
        # Handle APA italic markers (*text*)
        if style == 'apa' and '*' in text:
            _add_apa_formatted_run(p, text)
        else:
            run = p.add_run(text)
            set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小四')
    
    return formatted


def _add_apa_formatted_run(paragraph, text):
    """Parse APA text with *italic* markers and add formatted runs."""
    import re
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小四', italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小四')


def main():
    parser = argparse.ArgumentParser(description='Add reference list to .docx')
    parser.add_argument('--docx', required=True, help='Input .docx file path')
    parser.add_argument('--refs', required=True, help='JSON file with reference data')
    parser.add_argument('--style', default='gbt7714', choices=['gbt7714', 'apa'],
                        help='Citation style')
    parser.add_argument('--heading', default='参考文献', help='Section heading')
    parser.add_argument('--output', default='', help='Output path')
    
    args = parser.parse_args()
    output = args.output or args.docx
    
    doc = Document(args.docx)
    
    with open(args.refs, 'r', encoding='utf-8') as f:
        references = json.load(f)
    
    add_reference_list(doc, references, style=args.style, heading=args.heading)
    
    os.makedirs(os.path.dirname(output) or '.', exist_ok=True)
    doc.save(output)
    print(f"References added: {output}")


if __name__ == '__main__':
    main()
