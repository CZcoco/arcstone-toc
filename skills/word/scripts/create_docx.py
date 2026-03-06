"""
create_docx.py — Create a Chinese academic paper .docx with standard formatting.
Supports CLI and programmatic usage.
"""

import argparse
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from format_utils import (
    set_run_font, set_paragraph_format, add_body_paragraph,
    add_heading_paragraph, CN_BODY_FONT, CN_HEADING_FONT, EN_FONT,
    FONT_SIZE_MAP
)


def setup_page(doc, top=2.54, bottom=2.54, left=3.17, right=3.17):
    """Set A4 page size and margins (in cm)."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    return section


def setup_header_footer(doc, header_text="", show_page_number=True):
    """Add header text and page numbers."""
    section = doc.sections[0]
    
    # Header
    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        set_paragraph_format(hp, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = hp.add_run(header_text)
        set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小五')
        # Add bottom border to header
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    
    # Footer with page number
    if show_page_number:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        set_paragraph_format(fp, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = fp.add_run()
        set_run_font(run, en_font=EN_FONT, size='小五')
        # Insert PAGE field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar2)


def add_title(doc, title, font_size='二号'):
    """Add paper title."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=Pt(24), space_after=Pt(12))
    run = p.add_run(title)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size=font_size, bold=True)
    return p


def add_author_info(doc, author="", institution="", date_str=""):
    """Add author, institution, and date below title."""
    for text in [author, institution, date_str]:
        if text:
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 line_spacing=1.5, space_before=Pt(0), space_after=Pt(2))
            run = p.add_run(text)
            set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='四号')


def add_abstract(doc, abstract_text, keywords=None, lang='zh'):
    """Add abstract section with keywords."""
    # Abstract heading
    label = '摘  要' if lang == 'zh' else 'Abstract'
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=Pt(12), space_after=Pt(6))
    run = p.add_run(label)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='小四', bold=True)
    
    # Abstract body
    add_body_paragraph(doc, abstract_text, first_line_indent=True, font_size='小四')
    
    # Keywords
    if keywords:
        kw_label = '关键词：' if lang == 'zh' else 'Keywords: '
        kw_text = '；'.join(keywords) if lang == 'zh' else '; '.join(keywords)
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             line_spacing=1.5, space_before=Pt(6), space_after=Pt(6),
                             first_line_indent=Cm(0.74))
        run_label = p.add_run(kw_label)
        set_run_font(run_label, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='小四', bold=True)
        run_kw = p.add_run(kw_text)
        set_run_font(run_kw, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小四')


def create_academic_paper(title="论文标题", author="", institution="",
                          date_str="", abstract="", keywords=None,
                          header_text="", sections=None):
    """
    Create a complete academic paper document.
    
    Args:
        title: Paper title
        author: Author name(s)
        institution: Institution name
        date_str: Date string
        abstract: Abstract text
        keywords: List of keywords
        header_text: Header text (e.g., journal name)
        sections: List of dicts with 'heading', 'level', 'content' keys
    
    Returns:
        Document object
    """
    doc = Document()
    
    # Page setup
    setup_page(doc)
    setup_header_footer(doc, header_text=header_text)
    
    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), CN_BODY_FONT)
    style.paragraph_format.line_spacing = 1.5
    
    # Title page elements
    add_title(doc, title)
    add_author_info(doc, author, institution, date_str)
    
    # Abstract
    if abstract:
        add_abstract(doc, abstract, keywords, lang='zh')
    
    # Sections
    if sections:
        for sec in sections:
            heading = sec.get('heading', '')
            level = sec.get('level', 1)
            content = sec.get('content', '')
            add_heading_paragraph(doc, heading, level=level)
            if content:
                # Split by newlines for multiple paragraphs
                for para_text in content.split('\n'):
                    if para_text.strip():
                        add_body_paragraph(doc, para_text.strip())
    
    return doc


def main():
    parser = argparse.ArgumentParser(description='Create academic paper .docx')
    parser.add_argument('--title', default='论文标题', help='Paper title')
    parser.add_argument('--author', default='', help='Author name')
    parser.add_argument('--institution', default='', help='Institution')
    parser.add_argument('--date', default='', help='Date string')
    parser.add_argument('--abstract', default='', help='Abstract text')
    parser.add_argument('--keywords', nargs='*', help='Keywords list')
    parser.add_argument('--header', default='', help='Header text')
    parser.add_argument('--output', default='/workspace/paper.docx', help='Output path')
    
    args = parser.parse_args()
    
    doc = create_academic_paper(
        title=args.title,
        author=args.author,
        institution=args.institution,
        date_str=args.date,
        abstract=args.abstract,
        keywords=args.keywords,
        header_text=args.header
    )
    
    os.makedirs(os.path.dirname(args.output) or '.', exist_ok=True)
    doc.save(args.output)
    print(f"Document created: {args.output}")


if __name__ == '__main__':
    main()
