"""
add_formula.py — Convert LaTeX formulas to Word OMML native equations and insert into .docx.
Conversion path: LaTeX → MathML → OMML
"""

import argparse
import sys
import os
import re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

from format_utils import (
    set_run_font, set_paragraph_format, CN_BODY_FONT, EN_FONT
)

# XSLT for MathML to OMML conversion
_MATHML2OMML_XSLT = None


def _get_xslt():
    """Load the MathML-to-OMML XSLT stylesheet."""
    global _MATHML2OMML_XSLT
    if _MATHML2OMML_XSLT is not None:
        return _MATHML2OMML_XSLT
    
    # Try to find the XSLT file from common locations
    xslt_paths = [
        os.path.join(os.path.dirname(__file__), 'MML2OMML.XSL'),
        # Common Windows Office paths
        r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL',
        r'C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL',
        r'C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL',
        r'C:\Program Files\Microsoft Office\Office15\MML2OMML.XSL',
    ]
    
    for path in xslt_paths:
        if os.path.exists(path):
            with open(path, 'rb') as f:
                xslt_tree = etree.parse(f)
            _MATHML2OMML_XSLT = etree.XSLT(xslt_tree)
            return _MATHML2OMML_XSLT
    
    return None


def latex_to_mathml(latex_str):
    """Convert LaTeX string to MathML using latex2mathml."""
    import latex2mathml.converter
    mathml = latex2mathml.converter.convert(latex_str)
    return mathml


def mathml_to_omml(mathml_str):
    """Convert MathML string to OMML using XSLT."""
    xslt = _get_xslt()
    if xslt is None:
        return _mathml_to_omml_fallback(mathml_str)
    
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = xslt(mathml_tree)
    return omml_tree.getroot()


def _mathml_to_omml_fallback(mathml_str):
    """
    Fallback: Build OMML directly from LaTeX when XSLT is unavailable.
    Uses python-docx's oxml to construct oMath elements.
    """
    # Parse MathML and build a simplified OMML representation
    nsmap = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    
    # Create oMath element
    omath = etree.SubElement(
        etree.Element('dummy'),
        qn('m:oMath'),
        nsmap={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    )
    
    # For fallback, insert the MathML content as a run with the formula text
    # This is a simplified approach - the XSLT method is preferred
    return omath


def latex_to_omml(latex_str):
    """Convert LaTeX to OMML element, ready to insert into docx."""
    mathml_str = latex_to_mathml(latex_str)
    
    xslt = _get_xslt()
    if xslt is not None:
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omml_tree = xslt(mathml_tree)
        return omml_tree.getroot()
    else:
        # Fallback: use latex2mathml's output and wrap in oMathPara
        # We'll construct OMML manually from the MathML
        return _build_omml_from_mathml(mathml_str)


def _build_omml_from_mathml(mathml_str):
    """Build OMML from MathML by parsing the MathML tree and mapping elements."""
    M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Parse MathML
    try:
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    except Exception:
        # If parsing fails, create a simple text run
        omath = etree.Element(qn('m:oMath'))
        r = etree.SubElement(omath, qn('m:r'))
        t = etree.SubElement(r, qn('m:t'))
        t.text = mathml_str
        return omath
    
    omath = etree.Element(qn('m:oMath'))
    _convert_mathml_node(mathml_tree, omath)
    return omath


def _convert_mathml_node(node, parent):
    """Recursively convert MathML nodes to OMML."""
    tag = etree.QName(node.tag).localname if '}' in str(node.tag) else node.tag
    
    if tag == 'math':
        for child in node:
            _convert_mathml_node(child, parent)
    
    elif tag == 'mrow':
        for child in node:
            _convert_mathml_node(child, parent)
    
    elif tag == 'mi' or tag == 'mn' or tag == 'mo':
        r = etree.SubElement(parent, qn('m:r'))
        # Add run properties for italic (mi) or normal (mn, mo)
        if tag == 'mi' and node.text and len(node.text) == 1:
            rPr = etree.SubElement(r, qn('m:rPr'))
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'i')
        elif tag == 'mn' or tag == 'mo':
            rPr = etree.SubElement(r, qn('m:rPr'))
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'p')
        t = etree.SubElement(r, qn('m:t'))
        t.text = node.text or ''
    
    elif tag == 'msub':
        children = list(node)
        if len(children) >= 2:
            sSub = etree.SubElement(parent, qn('m:sSub'))
            e = etree.SubElement(sSub, qn('m:e'))
            _convert_mathml_node(children[0], e)
            sub = etree.SubElement(sSub, qn('m:sub'))
            _convert_mathml_node(children[1], sub)
    
    elif tag == 'msup':
        children = list(node)
        if len(children) >= 2:
            sSup = etree.SubElement(parent, qn('m:sSup'))
            e = etree.SubElement(sSup, qn('m:e'))
            _convert_mathml_node(children[0], e)
            sup = etree.SubElement(sSup, qn('m:sup'))
            _convert_mathml_node(children[1], sup)
    
    elif tag == 'msubsup':
        children = list(node)
        if len(children) >= 3:
            sSubSup = etree.SubElement(parent, qn('m:sSubSup'))
            e = etree.SubElement(sSubSup, qn('m:e'))
            _convert_mathml_node(children[0], e)
            sub = etree.SubElement(sSubSup, qn('m:sub'))
            _convert_mathml_node(children[1], sub)
            sup = etree.SubElement(sSubSup, qn('m:sup'))
            _convert_mathml_node(children[2], sup)
    
    elif tag == 'mfrac':
        children = list(node)
        if len(children) >= 2:
            f = etree.SubElement(parent, qn('m:f'))
            num = etree.SubElement(f, qn('m:num'))
            _convert_mathml_node(children[0], num)
            den = etree.SubElement(f, qn('m:den'))
            _convert_mathml_node(children[1], den)
    
    elif tag == 'msqrt':
        rad = etree.SubElement(parent, qn('m:rad'))
        radPr = etree.SubElement(rad, qn('m:radPr'))
        degHide = etree.SubElement(radPr, qn('m:degHide'))
        degHide.set(qn('m:val'), '1')
        deg = etree.SubElement(rad, qn('m:deg'))
        e = etree.SubElement(rad, qn('m:e'))
        for child in node:
            _convert_mathml_node(child, e)
    
    elif tag == 'mover':
        children = list(node)
        if len(children) >= 2:
            acc = etree.SubElement(parent, qn('m:acc'))
            accPr = etree.SubElement(acc, qn('m:accPr'))
            chr_elem = etree.SubElement(accPr, qn('m:chr'))
            # Get the accent character
            accent_text = children[1].text or '\u0302'
            chr_elem.set(qn('m:val'), accent_text)
            e = etree.SubElement(acc, qn('m:e'))
            _convert_mathml_node(children[0], e)
    
    elif tag == 'munder':
        children = list(node)
        if len(children) >= 2:
            # Use groupChr for underscripts
            groupChr = etree.SubElement(parent, qn('m:groupChr'))
            groupChrPr = etree.SubElement(groupChr, qn('m:groupChrPr'))
            pos = etree.SubElement(groupChrPr, qn('m:pos'))
            pos.set(qn('m:val'), 'bot')
            e = etree.SubElement(groupChr, qn('m:e'))
            _convert_mathml_node(children[0], e)
    
    elif tag == 'mtext':
        r = etree.SubElement(parent, qn('m:r'))
        rPr = etree.SubElement(r, qn('m:rPr'))
        sty = etree.SubElement(rPr, qn('m:sty'))
        sty.set(qn('m:val'), 'p')
        t = etree.SubElement(r, qn('m:t'))
        t.text = node.text or ''
    
    elif tag == 'mspace':
        r = etree.SubElement(parent, qn('m:r'))
        t = etree.SubElement(r, qn('m:t'))
        t.text = ' '
    
    else:
        # Generic fallback: process children
        for child in node:
            _convert_mathml_node(child, parent)
        if node.text:
            r = etree.SubElement(parent, qn('m:r'))
            t = etree.SubElement(r, qn('m:t'))
            t.text = node.text


def insert_formula(doc, latex_str, label="", centered=True):
    """
    Insert a LaTeX formula as a native Word equation.
    
    Args:
        doc: Document object
        latex_str: LaTeX formula string (without $ delimiters)
        label: Equation label, e.g. "(1)"
        centered: Whether to center the formula (default True)
    
    Returns:
        The paragraph containing the formula
    """
    # Create paragraph
    p = doc.add_paragraph()
    
    if label:
        # Use a table-like layout: formula centered, label right-aligned
        # We use tab stops for this
        set_paragraph_format(p, line_spacing=1.5,
                             space_before=Pt(6), space_after=Pt(6))
        pPr = p._p.get_or_add_pPr()
        
        # Set tab stops: center tab at page center, right tab at right margin
        tabs = parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="center" w:pos="4153"/>'
            f'  <w:tab w:val="right" w:pos="8306"/>'
            f'</w:tabs>'
        )
        pPr.append(tabs)
        
        # Add center tab
        run_tab1 = p.add_run()
        tab1 = parse_xml(f'<w:tab {nsdecls("w")}/>')
        run_tab1._r.append(tab1)
        
        # Add formula
        omml = latex_to_omml(latex_str)
        p._p.append(omml)
        
        # Add right tab + label
        run_tab2 = p.add_run()
        tab2 = parse_xml(f'<w:tab {nsdecls("w")}/>')
        run_tab2._r.append(tab2)
        run_label = p.add_run(label)
        set_run_font(run_label, en_font=EN_FONT, size='小四')
    else:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=1.5, space_before=Pt(6), space_after=Pt(6))
        omml = latex_to_omml(latex_str)
        p._p.append(omml)
    
    return p


def main():
    parser = argparse.ArgumentParser(description='Insert LaTeX formula into .docx as OMML')
    parser.add_argument('--docx', required=True, help='Input .docx file path')
    parser.add_argument('--latex', required=True, help='LaTeX formula string')
    parser.add_argument('--label', default='', help='Equation label, e.g. "(1)"')
    parser.add_argument('--output', default='', help='Output path (default: overwrite input)')
    
    args = parser.parse_args()
    output = args.output or args.docx
    
    doc = Document(args.docx)
    insert_formula(doc, args.latex, label=args.label)
    
    os.makedirs(os.path.dirname(output) or '.', exist_ok=True)
    doc.save(output)
    print(f"Formula inserted: {output}")


if __name__ == '__main__':
    main()
