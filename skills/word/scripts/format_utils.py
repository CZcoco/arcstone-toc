"""
format_utils.py — Shared utilities for Word academic paper formatting.
Provides unit conversions, font helpers, and paragraph formatting functions.
"""

from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import os

# ── Chinese Academic Font Size Mapping ──
FONT_SIZE_MAP = {
    '初号': Pt(42),    '小初': Pt(36),
    '一号': Pt(26),    '小一': Pt(24),
    '二号': Pt(22),    '小二': Pt(18),
    '三号': Pt(16),    '小三': Pt(15),
    '四号': Pt(14),    '小四': Pt(12),
    '五号': Pt(10.5),  '小五': Pt(9),
    '六号': Pt(7.5),   '小六': Pt(6.5),
    '七号': Pt(5.5),   '八号': Pt(5),
}

# ── Default Academic Fonts ──
CN_BODY_FONT = '宋体'        # SimSun
CN_HEADING_FONT = '黑体'     # SimHei
EN_FONT = 'Times New Roman'

# ── Fallback fonts if Chinese fonts unavailable ──
CN_BODY_FALLBACK = 'SimSun'
CN_HEADING_FALLBACK = 'SimHei'


def set_run_font(run, cn_font=None, en_font=EN_FONT, size=None, bold=False, italic=False, color=None):
    """Set font properties for a run, handling both Chinese and English fonts."""
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size if isinstance(size, (int, Emu)) else FONT_SIZE_MAP.get(size, Pt(12))
    if en_font:
        run.font.name = en_font
    if cn_font:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    if color:
        run.font.color.rgb = color if isinstance(color, RGBColor) else RGBColor.from_string(color)


def set_paragraph_format(paragraph, alignment=None, line_spacing=None,
                         space_before=None, space_after=None,
                         first_line_indent=None, left_indent=None):
    """Set paragraph formatting properties."""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if line_spacing is not None:
        if isinstance(line_spacing, (int, float)):
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = line_spacing
        else:
            pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent


def add_body_paragraph(doc, text, first_line_indent=True, font_size='小四',
                       cn_font=CN_BODY_FONT, en_font=EN_FONT, line_spacing=1.5):
    """Add a standard body paragraph with academic formatting."""
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=line_spacing,
        space_before=Pt(0),
        space_after=Pt(0),
        first_line_indent=Cm(0.74) if first_line_indent else None  # 2 chars indent
    )
    run = p.add_run(text)
    set_run_font(run, cn_font=cn_font, en_font=en_font, size=font_size)
    return p


def add_heading_paragraph(doc, text, level=1):
    """Add a heading with proper Chinese academic formatting."""
    size_map = {1: '三号', 2: '四号', 3: '小四'}
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_before=Pt(12),
        space_after=Pt(6)
    )
    run = p.add_run(text)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT,
                 size=size_map.get(level, '小四'), bold=True)
    return p


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a table cell. Each border param is a dict like {'sz': 12, 'val': 'single', 'color': '000000'}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, props in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if props:
            tag = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{props.get("val", "single")}" '
                f'w:sz="{props.get("sz", 4)}" w:space="0" '
                f'w:color="{props.get("color", "000000")}"/>'
            )
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def remove_all_borders(table):
    """Remove all borders from a table (used before applying three-line style)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def apply_three_line_style(table, header_rows=1):
    """Apply three-line table style: top 1.5pt, header-bottom 0.75pt, bottom 1.5pt."""
    remove_all_borders(table)
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    
    top_border = {'sz': 12, 'val': 'single', 'color': '000000'}      # 1.5pt
    mid_border = {'sz': 6, 'val': 'single', 'color': '000000'}       # 0.75pt
    bot_border = {'sz': 12, 'val': 'single', 'color': '000000'}      # 1.5pt
    
    for col_idx in range(num_cols):
        # Top border on first row
        set_cell_border(table.cell(0, col_idx), top=top_border)
        # Mid border: bottom of header row(s)
        set_cell_border(table.cell(header_rows - 1, col_idx), bottom=mid_border)
        # Bottom border on last row
        set_cell_border(table.cell(num_rows - 1, col_idx), bottom=bot_border)


def format_table_text(cell, text, cn_font=CN_BODY_FONT, en_font=EN_FONT,
                      size='五号', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=False):
    """Format text in a table cell."""
    cell.text = ''
    p = cell.paragraphs[0]
    set_paragraph_format(p, alignment=alignment, line_spacing=1.0,
                         space_before=Pt(2), space_after=Pt(2))
    run = p.add_run(str(text))
    set_run_font(run, cn_font=cn_font, en_font=en_font, size=size, bold=bold)


# ═══════════════════════════════════════════════════════════════
# ── TOC (Table of Contents) ──
# ═══════════════════════════════════════════════════════════════

def setup_heading_styles(doc):
    """
    Configure built-in Heading 1/2/3 styles with Chinese academic formatting.
    This is required for TOC to work — Word TOC relies on built-in Heading styles.
    """
    heading_config = {
        'Heading 1': {'cn': CN_HEADING_FONT, 'en': EN_FONT, 'size': Pt(16), 'bold': True,
                      'align': WD_ALIGN_PARAGRAPH.CENTER, 'before': Pt(12), 'after': Pt(6)},
        'Heading 2': {'cn': CN_HEADING_FONT, 'en': EN_FONT, 'size': Pt(14), 'bold': True,
                      'align': WD_ALIGN_PARAGRAPH.LEFT, 'before': Pt(12), 'after': Pt(6)},
        'Heading 3': {'cn': CN_HEADING_FONT, 'en': EN_FONT, 'size': Pt(12), 'bold': True,
                      'align': WD_ALIGN_PARAGRAPH.LEFT, 'before': Pt(6), 'after': Pt(4)},
    }
    for style_name, cfg in heading_config.items():
        style = doc.styles[style_name]
        style.font.name = cfg['en']
        style.font.size = cfg['size']
        style.font.bold = cfg['bold']
        style._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['cn'])
        style.paragraph_format.alignment = cfg['align']
        style.paragraph_format.space_before = cfg['before']
        style.paragraph_format.space_after = cfg['after']
        style.paragraph_format.line_spacing = 1.5


def add_heading_with_style(doc, text, level=1):
    """
    Add a heading using built-in Heading styles (required for TOC).
    Use this instead of add_heading_paragraph when TOC is needed.
    """
    p = doc.add_heading(text, level=level)
    # Override font since add_heading uses default style
    for run in p.runs:
        set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT,
                     size={1: '三号', 2: '四号', 3: '小四'}.get(level, '小四'), bold=True)
    return p


def insert_toc(doc, title="目  录", max_level=3):
    """
    Insert a Table of Contents with a TOC field code.
    The TOC will auto-update when opened in Word (right-click → Update Field).
    
    Args:
        doc: Document object
        title: TOC heading text
        max_level: Maximum heading level to include (1-3)
    
    Returns:
        The TOC paragraph
    """
    # TOC title
    p_title = doc.add_paragraph()
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=Pt(24), space_after=Pt(18))
    run = p_title.add_run(title)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='三号', bold=True)

    # TOC field
    p_toc = doc.add_paragraph()
    run = p_toc.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar_begin)

    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-{max_level}" \\h \\z \\u </w:instrText>'
    )
    run._r.append(instr)

    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar_separate)

    # Placeholder text (replaced when user updates field in Word)
    run_placeholder = p_toc.add_run("（请在 Word 中右键此处 → 更新域，即可生成目录）")
    set_run_font(run_placeholder, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小四')

    run_end = p_toc.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end._r.append(fldChar_end)

    return p_toc


# ═══════════════════════════════════════════════════════════════
# ── Figure / Image Insertion ──
# ═══════════════════════════════════════════════════════════════

def insert_image(doc, image_path, width_cm=14, caption="", caption_position='below'):
    """
    Insert an image with a centered caption (图题).
    
    Args:
        doc: Document object
        image_path: Path to image file (.png, .jpg, etc.)
        width_cm: Image width in cm (default 14, fits A4 with margins)
        caption: Caption text, e.g. "图1 变量趋势图"
        caption_position: 'below' (default, Chinese standard) or 'above'
    
    Returns:
        Tuple of (image_paragraph, caption_paragraph) or (image_paragraph, None)
    """
    cap_p = None

    # Caption above (rare, but supported)
    if caption and caption_position == 'above':
        cap_p = _add_figure_caption(doc, caption)

    # Image paragraph
    p_img = doc.add_paragraph()
    set_paragraph_format(p_img, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.0, space_before=Pt(6), space_after=Pt(2))
    run = p_img.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    # Caption below (Chinese standard)
    if caption and caption_position == 'below':
        cap_p = _add_figure_caption(doc, caption)

    return p_img, cap_p


def _add_figure_caption(doc, caption):
    """Add a figure caption paragraph."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.0, space_before=Pt(2), space_after=Pt(12))
    run = p.add_run(caption)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='五号', bold=True)
    return p


# ═══════════════════════════════════════════════════════════════
# ── Footnotes ──
# ═══════════════════════════════════════════════════════════════

def add_footnote(paragraph, footnote_text):
    """
    Add a footnote to a paragraph. Inserts a superscript reference number
    at the current end of the paragraph and creates the footnote content.
    
    Args:
        paragraph: The paragraph to attach the footnote to
        footnote_text: The footnote content text
    
    Returns:
        The footnote element
    """
    # Get or create footnotes part
    doc_part = paragraph.part
    
    # Access the footnotes XML
    # python-docx doesn't natively support footnotes, so we use oxml directly
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    # Find or create footnotes part
    footnotes_part = None
    for rel in doc_part.rels.values():
        if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes':
            footnotes_part = rel.target_part
            break
    
    if footnotes_part is None:
        # Create footnotes part from scratch
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        
        footnotes_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        footnotes_part = Part(
            PackURI('/word/footnotes.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
            footnotes_xml.encode('utf-8'),
            doc_part.package
        )
        doc_part.relate_to(
            footnotes_part,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        )
    
    # Parse existing footnotes to find next ID
    from lxml import etree
    footnotes_elem = etree.fromstring(footnotes_part.blob)
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    existing_ids = [int(fn.get(qn('w:id'))) for fn in footnotes_elem.findall(qn('w:footnote'))
                    if fn.get(qn('w:id')) not in ('-1', '0')]
    next_id = max(existing_ids, default=0) + 1
    
    # Create footnote element
    footnote_xml = (
        f'<w:footnote xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:id="{next_id}">'
        f'<w:p>'
        f'<w:pPr><w:pStyle w:val="FootnoteText"/>'
        f'<w:rPr><w:sz w:val="18"/><w:szCs w:val="18"/></w:rPr>'
        f'</w:pPr>'
        f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteRef/></w:r>'
        f'<w:r><w:t xml:space="preserve"> {footnote_text}</w:t></w:r>'
        f'</w:p>'
        f'</w:footnote>'
    )
    new_fn = etree.fromstring(footnote_xml)
    footnotes_elem.append(new_fn)
    footnotes_part._blob = etree.tostring(footnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True)
    
    # Insert footnote reference in the paragraph
    run = paragraph.add_run()
    rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rStyle w:val="FootnoteReference"/>'
        f'<w:vertAlign w:val="superscript"/>'
        f'</w:rPr>'
    )
    run._r.insert(0, rPr)
    footnoteRef = parse_xml(
        f'<w:footnoteReference {nsdecls("w")} w:id="{next_id}"/>'
    )
    run._r.append(footnoteRef)
    
    return new_fn


# ═══════════════════════════════════════════════════════════════
# ── Cover Page (毕业论文封面) ──
# ═══════════════════════════════════════════════════════════════

def add_cover_page(doc, title="论文标题", author="", student_id="",
                   major="", supervisor="", institution="",
                   degree="硕士", date_str="", logo_path=None):
    """
    Add a thesis cover page with standard layout.
    
    Args:
        doc: Document object
        title: Thesis title
        author: Author name
        student_id: Student ID number
        major: Major/discipline
        supervisor: Supervisor name
        institution: University/institution name
        degree: Degree type (硕士/博士/本科)
        date_str: Date string, e.g. "2025年6月"
        logo_path: Optional path to university logo image
    
    Returns:
        None (modifies doc in place)
    """
    from docx.shared import Pt, Cm

    # University logo (if provided)
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        set_paragraph_format(p_logo, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(36), space_after=Pt(12))
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Cm(5))

    # Institution name
    if institution:
        p_inst = doc.add_paragraph()
        set_paragraph_format(p_inst, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(24) if not logo_path else Pt(6),
                             space_after=Pt(6))
        run = p_inst.add_run(institution)
        set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='小初', bold=True)

    # Degree type
    degree_text = f"{degree}学位论文"
    p_deg = doc.add_paragraph()
    set_paragraph_format(p_deg, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(12), space_after=Pt(36))
    run = p_deg.add_run(degree_text)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='一号', bold=True)

    # Title
    p_title = doc.add_paragraph()
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=Pt(24), space_after=Pt(48))
    run = p_title.add_run(title)
    set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='二号', bold=True)

    # Info fields (underlined style)
    fields = []
    if author:
        fields.append(('姓    名', author))
    if student_id:
        fields.append(('学    号', student_id))
    if major:
        fields.append(('专    业', major))
    if supervisor:
        fields.append(('指导教师', supervisor))

    for label, value in fields:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=2.0, space_before=Pt(2), space_after=Pt(2))
        # Label
        run_label = p.add_run(f"{label}：")
        set_run_font(run_label, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='三号')
        # Value with underline
        run_val = p.add_run(f"  {value}  ")
        set_run_font(run_val, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='三号')
        run_val.underline = True

    # Date
    if date_str:
        p_date = doc.add_paragraph()
        set_paragraph_format(p_date, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(48), space_after=Pt(12))
        run = p_date.add_run(date_str)
        set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='三号')

    # Page break after cover
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# ── Page Break / Section Break ──
# ═══════════════════════════════════════════════════════════════

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def add_section_break(doc, start_type='new_page'):
    """
    Add a section break.
    
    Args:
        start_type: 'new_page', 'continuous', 'even_page', 'odd_page'
    """
    from docx.enum.section import WD_SECTION_START
    type_map = {
        'new_page': WD_SECTION_START.NEW_PAGE,
        'continuous': WD_SECTION_START.CONTINUOUS,
        'even_page': WD_SECTION_START.EVEN_PAGE,
        'odd_page': WD_SECTION_START.ODD_PAGE,
    }
    new_section = doc.add_section(type_map.get(start_type, WD_SECTION_START.NEW_PAGE))
    return new_section
