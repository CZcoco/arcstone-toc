"""
modify_docx.py — Read and modify existing .docx documents.
Reformat fonts, margins, headings, line spacing, and other properties.
Supports both CLI and programmatic usage.
"""

import argparse
import sys
import os
import re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from format_utils import (
    set_run_font, set_paragraph_format, setup_heading_styles,
    CN_BODY_FONT, CN_HEADING_FONT, EN_FONT, FONT_SIZE_MAP
)


def load_docx(path):
    """Load an existing .docx file."""
    return Document(path)


def reformat_all_fonts(doc, cn_body=CN_BODY_FONT, cn_heading=CN_HEADING_FONT,
                       en_font=EN_FONT, body_size='小四', heading_sizes=None):
    """
    Reformat all fonts in the document to academic standard.
    
    Args:
        doc: Document object
        cn_body: Chinese body font (default 宋体)
        cn_heading: Chinese heading font (default 黑体)
        en_font: English font (default Times New Roman)
        body_size: Body text size (default 小四)
        heading_sizes: Dict mapping heading level to size, e.g. {1: '三号', 2: '四号'}
    """
    if heading_sizes is None:
        heading_sizes = {1: '三号', 2: '四号', 3: '小四'}
    
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        is_heading = style_name.startswith('Heading')
        heading_level = 0
        if is_heading:
            try:
                heading_level = int(style_name.replace('Heading ', ''))
            except ValueError:
                heading_level = 1
        
        for run in paragraph.runs:
            if is_heading:
                size = heading_sizes.get(heading_level, '小四')
                set_run_font(run, cn_font=cn_heading, en_font=en_font, size=size, bold=True)
            else:
                set_run_font(run, cn_font=cn_body, en_font=en_font, size=body_size)


def reformat_margins(doc, top=2.54, bottom=2.54, left=3.17, right=3.17):
    """Reformat page margins for all sections (in cm)."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def reformat_line_spacing(doc, body_spacing=1.5, heading_spacing=1.5):
    """Reformat line spacing for all paragraphs."""
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        is_heading = style_name.startswith('Heading')
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = heading_spacing if is_heading else body_spacing


def reformat_first_line_indent(doc, indent_cm=0.74, skip_headings=True):
    """Add first-line indent to body paragraphs."""
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        is_heading = style_name.startswith('Heading')
        if skip_headings and is_heading:
            continue
        # Skip empty paragraphs and centered paragraphs (titles, captions)
        if not paragraph.text.strip():
            continue
        if paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue
        paragraph.paragraph_format.first_line_indent = Cm(indent_cm)


def reformat_page_size(doc, width_cm=21.0, height_cm=29.7):
    """Set page size for all sections (default A4)."""
    for section in doc.sections:
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)


def reformat_paragraph_spacing(doc, space_before=Pt(0), space_after=Pt(0),
                                heading_before=Pt(12), heading_after=Pt(6)):
    """Reformat paragraph spacing."""
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        is_heading = style_name.startswith('Heading')
        pf = paragraph.paragraph_format
        if is_heading:
            pf.space_before = heading_before
            pf.space_after = heading_after
        else:
            pf.space_before = space_before
            pf.space_after = space_after


def reformat_tables_to_three_line(doc):
    """Convert all tables in the document to three-line style."""
    from format_utils import apply_three_line_style, format_table_text
    for table in doc.tables:
        apply_three_line_style(table, header_rows=1)
        # Reformat table text fonts
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='五号')


def apply_academic_format(doc, margins=True, fonts=True, spacing=True,
                          indent=True, page_size=True, tables=False):
    """
    One-click: apply full Chinese academic formatting to an existing document.
    
    Args:
        doc: Document object
        margins: Reformat margins to standard
        fonts: Reformat all fonts to academic standard
        spacing: Reformat line spacing
        indent: Add first-line indent
        page_size: Set A4 page size
        tables: Convert tables to three-line style
    """
    if page_size:
        reformat_page_size(doc)
    if margins:
        reformat_margins(doc)
    if fonts:
        setup_heading_styles(doc)
        reformat_all_fonts(doc)
    if spacing:
        reformat_line_spacing(doc)
        reformat_paragraph_spacing(doc)
    if indent:
        reformat_first_line_indent(doc)
    if tables:
        reformat_tables_to_three_line(doc)


def get_document_info(doc):
    """Get summary info about a document."""
    info = {
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'sections': len(doc.sections),
        'headings': [],
        'styles_used': set(),
    }
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else 'None'
        info['styles_used'].add(style_name)
        if style_name.startswith('Heading'):
            info['headings'].append({'level': style_name, 'text': p.text[:50]})
    info['styles_used'] = sorted(info['styles_used'])
    return info


def main():
    parser = argparse.ArgumentParser(description='Modify existing .docx formatting')
    parser.add_argument('input', help='Input .docx file path')
    parser.add_argument('--output', default='', help='Output path (default: overwrite)')
    parser.add_argument('--info', action='store_true', help='Print document info and exit')
    parser.add_argument('--academic', action='store_true', help='Apply full academic formatting')
    parser.add_argument('--fonts', action='store_true', help='Reformat fonts only')
    parser.add_argument('--margins', action='store_true', help='Reformat margins only')
    parser.add_argument('--spacing', action='store_true', help='Reformat line spacing only')
    parser.add_argument('--indent', action='store_true', help='Add first-line indent only')
    parser.add_argument('--tables', action='store_true', help='Convert tables to three-line')
    
    args = parser.parse_args()
    output = args.output or args.input
    
    doc = load_docx(args.input)
    
    if args.info:
        import json
        info = get_document_info(doc)
        print(json.dumps(info, ensure_ascii=False, indent=2))
        return
    
    if args.academic:
        apply_academic_format(doc, tables=args.tables)
        print(f"Academic formatting applied: {output}")
    else:
        if args.fonts:
            reformat_all_fonts(doc)
            print("Fonts reformatted.")
        if args.margins:
            reformat_margins(doc)
            print("Margins reformatted.")
        if args.spacing:
            reformat_line_spacing(doc)
            reformat_paragraph_spacing(doc)
            print("Spacing reformatted.")
        if args.indent:
            reformat_first_line_indent(doc)
            print("First-line indent applied.")
        if args.tables:
            reformat_tables_to_three_line(doc)
            print("Tables converted to three-line style.")
    
    os.makedirs(os.path.dirname(output) or '.', exist_ok=True)
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == '__main__':
    main()
