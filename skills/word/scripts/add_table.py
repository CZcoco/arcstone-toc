"""
add_table.py — Insert academic three-line tables into .docx documents.
Supports CSV input or programmatic list data.
"""

import argparse
import csv
import json
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from format_utils import (
    set_run_font, set_paragraph_format, apply_three_line_style,
    format_table_text, CN_BODY_FONT, CN_HEADING_FONT, EN_FONT
)


def insert_three_line_table(doc, headers, data, title="", note="",
                            col_widths=None, header_rows=1):
    """
    Insert a three-line table into the document.
    
    Args:
        doc: Document object
        headers: List of header strings, e.g. ["变量", "(1) OLS", "(2) FE"]
        data: List of lists, each inner list is a row
        title: Table title (displayed above table)
        note: Table note (displayed below table)
        col_widths: Optional list of column widths in cm
        header_rows: Number of header rows (default 1)
    
    Returns:
        The table object
    """
    # Table title
    if title:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=1.0, space_before=Pt(12), space_after=Pt(6))
        run = p.add_run(title)
        set_run_font(run, cn_font=CN_HEADING_FONT, en_font=EN_FONT, size='小四', bold=True)
    
    # Create table
    num_cols = len(headers)
    num_rows = len(data) + header_rows
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    # Fill header
    for j, h in enumerate(headers):
        format_table_text(table.cell(0, j), h, cn_font=CN_HEADING_FONT,
                          en_font=EN_FONT, size='五号', bold=True)
    
    # Fill data rows
    for i, row_data in enumerate(data):
        for j, cell_val in enumerate(row_data):
            if j < num_cols:
                format_table_text(table.cell(i + header_rows, j), cell_val,
                                  cn_font=CN_BODY_FONT, en_font=EN_FONT, size='五号')
    
    # Apply three-line style
    apply_three_line_style(table, header_rows=header_rows)
    
    # Table note
    if note:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.0, space_before=Pt(2), space_after=Pt(12))
        run = p.add_run(note)
        set_run_font(run, cn_font=CN_BODY_FONT, en_font=EN_FONT, size='小五')
    
    return table


def load_csv_data(csv_path):
    """Load headers and data from a CSV file."""
    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        return [], []
    headers = rows[0]
    data = rows[1:]
    return headers, data


def main():
    parser = argparse.ArgumentParser(description='Insert three-line table into .docx')
    parser.add_argument('--docx', required=True, help='Input .docx file path')
    parser.add_argument('--csv', default='', help='CSV file with table data')
    parser.add_argument('--json', default='', help='JSON file with {headers:[], data:[[]]}')
    parser.add_argument('--title', default='', help='Table title')
    parser.add_argument('--note', default='', help='Table note')
    parser.add_argument('--output', default='', help='Output path (default: overwrite input)')
    
    args = parser.parse_args()
    output = args.output or args.docx
    
    # Load document
    doc = Document(args.docx)
    
    # Load data
    if args.csv:
        headers, data = load_csv_data(args.csv)
    elif args.json:
        with open(args.json, 'r', encoding='utf-8') as f:
            j = json.load(f)
        headers = j['headers']
        data = j['data']
    else:
        print("Error: provide --csv or --json", file=sys.stderr)
        sys.exit(1)
    
    insert_three_line_table(doc, headers, data, title=args.title, note=args.note)
    
    os.makedirs(os.path.dirname(output) or '.', exist_ok=True)
    doc.save(output)
    print(f"Table inserted: {output}")


if __name__ == '__main__':
    main()
